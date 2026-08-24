import re
from io import BytesIO

from bs4 import BeautifulSoup


class ExtractionError(ValueError):
    pass


def extract_text_from_bytes(data: bytes, filename: str) -> str:
    """Extract plain text from an uploaded file based on its extension."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    try:
        if ext == "pdf":
            return _extract_pdf(data)
        if ext == "docx":
            return _extract_docx(data)
        if ext in ("html", "htm"):
            return _extract_html(data)
        if ext in ("txt", "md", "markdown"):
            return data.decode("utf-8", errors="replace")
    except Exception as exc:
        raise ExtractionError(f"Could not extract text from {filename}: {exc}") from exc
    raise ExtractionError(f"Unsupported file type: {ext or 'unknown'}")


def extract_text_from_html(data: str) -> str:
    """Extract readable text from raw HTML content (used for URL ingestion)."""
    soup = BeautifulSoup(data, "html.parser")
    for tag in soup(["script", "style", "noscript", "header", "footer", "nav"]):
        tag.decompose()
    text = soup.get_text(separator=" ")
    text = re.sub(r"[ \t]+", " ", text)
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    return "\n".join(lines)


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(data))
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n".join(pages).strip()


def _extract_docx(data: bytes) -> str:
    import docx

    document = docx.Document(BytesIO(data))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(c for c in cells if c))
    return "\n".join(parts).strip()


def _extract_html(data: bytes) -> str:
    return extract_text_from_html(data.decode("utf-8", errors="replace"))
